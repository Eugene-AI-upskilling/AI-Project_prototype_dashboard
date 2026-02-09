# coding=utf-8
"""
================================================================================
4_Earnings_Call_Summarizer.py
================================================================================
실적발표 컨퍼런스콜 요약 스크립트

[기능]
- 컨콜 원문(txt, docx, 직접 입력)을 읽어 구조화된 요약문 생성
- OpenAI GPT를 활용한 자동 요약
- 표준 양식으로 출력 (연간실적, 분기실적, 매출구성, 비용구조, 파이프라인, Q&A 등)
- docx 파일로 저장
- 텔레그램 발송 (선택)

[환경변수]
- OPENAI_API: OpenAI API 키
- BOT_TOKEN: 텔레그램 봇 토큰 (선택)
- CHAT_ID: 텔레그램 채팅 ID (선택)

[출력]
- output/earnings_call_summaries/{회사명}_{날짜}_컨콜요약.docx

[실행 방법]
$ python scripts/4_Earnings_Call_Summarizer.py --file="원문파일.txt"
$ python scripts/4_Earnings_Call_Summarizer.py --file="원문파일.docx" --telegram
$ python scripts/4_Earnings_Call_Summarizer.py --interactive

================================================================================
"""

import os
import sys
import re
import argparse
from datetime import datetime
from typing import Optional, Tuple
import warnings

warnings.filterwarnings('ignore')

# 환경변수 로드
_script_dir = os.path.dirname(os.path.abspath(__file__))
_project_dir = os.path.dirname(_script_dir)
_env_path = os.path.join(_project_dir, '.env')

from dotenv import load_dotenv
load_dotenv(dotenv_path=_env_path)

# API 키
OPENAI_API_KEY = os.getenv('OPENAI_API') or os.getenv('OPENAI_API_KEY')
BOT_TOKEN = os.getenv('BOT_TOKEN')
CHAT_ID = os.getenv('CHAT_ID')

# 출력 디렉토리
OUTPUT_DIR = os.path.join(_project_dir, 'output', 'earnings_call_summaries')


# =============================================================================
# 요약 프롬프트 템플릿
# =============================================================================

SUMMARY_PROMPT = """당신은 증권사 리서치 애널리스트입니다. 아래 실적발표 컨퍼런스콜 원문을 분석하여 정해진 양식에 맞게 요약해주세요.

## 출력 양식

< {회사명} {분기} 영업실적 콘퍼런스 콜 정리 >

# 연간 영업 실적
[매출] 금액 (YoY 변화율)
[EBITDA] 금액 (YoY 변화율) 또는 [영업이익] 금액
- 주요 특이사항 불릿포인트

# {분기} 영업 실적
[매출] 금액 (QoQ, YoY 변화율)
[EBITDA] 금액 (QoQ, YoY 변화율) 또는 [영업이익] 금액
[영업이익] 금액
[당기순이익] 금액
- 주요 특이사항 불릿포인트

# 매출 포트폴리오
주요 제품/서비스별 매출 비중 또는 금액
: 제품A X%, 제품B X%, ...

# 지역/부문별 매출 구성
[지역별] 국가/지역별 비중
[부문별] 사업부문별 비중 (해당시)

# 주요 비용구조
주요 비용 항목별 금액 및 증감
- 비용항목 금액 (증감률): 설명

# 향후 계획/파이프라인
[시기] 계획 내용
- 신제품, 신규 사업, 전략 방향 등

# 주주환원 정책
배당, 자사주 매입 등 주주환원 관련 내용

* Comment
- 컨콜 내용을 바탕으로 한 핵심 시사점 2-3개

Q&A
Q1. (질문자/기관)
Q1-1. 질문 내용?
- 답변 요약

Q2. (질문자/기관)
...

## 작성 지침
1. 금액은 억원 또는 조원 단위로 표기
2. 변화율은 YoY(전년동기비), QoQ(전분기비)로 표기
3. 구체적인 수치가 있으면 반드시 포함
4. Q&A는 핵심 질문과 답변만 간결하게 정리
5. 언급되지 않은 항목은 "(언급 없음)" 표시
6. 회사명과 분기는 원문에서 추출

## 컨퍼런스콜 원문:
{transcript}
"""


# =============================================================================
# 파일 읽기
# =============================================================================

def read_transcript_file(filepath: str) -> str:
    """
    컨콜 원문 파일 읽기

    지원 형식: .txt, .docx
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"파일을 찾을 수 없습니다: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.txt':
        # 여러 인코딩 시도
        for encoding in ['utf-8', 'cp949', 'euc-kr', 'utf-16']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"파일 인코딩을 인식할 수 없습니다: {filepath}")

    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)
        except ImportError:
            raise ImportError("python-docx 패키지가 필요합니다: pip install python-docx")

    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext} (지원: .txt, .docx)")


def extract_company_and_quarter(transcript: str) -> Tuple[str, str]:
    """원문에서 회사명과 분기 추출 시도"""
    company = "회사명"
    quarter = "분기"

    # 일반적인 패턴 매칭
    # 예: "삼성전자 4Q25", "넷마블 2025년 4분기"
    patterns = [
        r'([가-힣A-Za-z]+)\s*(\d{1,2}Q\d{2})',  # 넷마블 4Q25
        r'([가-힣A-Za-z]+)\s*(\d{4}년?\s*\d분기)',  # 넷마블 2025년 4분기
        r'([가-힣A-Za-z]+)\s*실적',  # XX 실적
    ]

    for pattern in patterns:
        match = re.search(pattern, transcript[:500])
        if match:
            company = match.group(1)
            if len(match.groups()) > 1:
                quarter = match.group(2)
            break

    return company, quarter


# =============================================================================
# GPT 요약
# =============================================================================

def summarize_with_gpt(transcript: str, model: str = "gpt-4o") -> str:
    """
    OpenAI GPT를 사용하여 컨콜 요약
    """
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API 환경변수가 설정되지 않았습니다.")

    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai 패키지가 필요합니다: pip install openai")

    client = OpenAI(api_key=OPENAI_API_KEY)

    # 회사명/분기 추출
    company, quarter = extract_company_and_quarter(transcript)

    # 프롬프트 구성
    prompt = SUMMARY_PROMPT.format(
        회사명=company,
        분기=quarter,
        transcript=transcript
    )

    print(f"  GPT 요약 중... (모델: {model})")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "당신은 증권사 리서치 애널리스트입니다. 실적발표 컨퍼런스콜을 정확하고 구조화된 형식으로 요약합니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=4000
    )

    summary = response.choices[0].message.content
    return summary, company, quarter


# =============================================================================
# 파일 저장
# =============================================================================

def save_to_docx(summary: str, company: str, quarter: str, output_dir: Optional[str] = None) -> str:
    """요약을 docx 파일로 저장"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx 패키지가 필요합니다: pip install python-docx")

    if output_dir is None:
        output_dir = OUTPUT_DIR

    os.makedirs(output_dir, exist_ok=True)

    # 파일명 생성
    date_str = datetime.now().strftime('%Y%m%d')
    filename = f"{company}_{date_str}_{quarter}_컨콜요약.docx"
    filepath = os.path.join(output_dir, filename)

    # docx 생성
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 내용 추가
    for line in summary.split('\n'):
        if line.startswith('< ') or line.startswith('# '):
            # 제목
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12) if line.startswith('< ') else Pt(11)
        elif line.startswith('* ') or line.startswith('Q'):
            # 코멘트/Q&A
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)

    doc.save(filepath)
    print(f"[저장 완료] {filepath}")

    return filepath


def save_to_txt(summary: str, company: str, quarter: str, output_dir: Optional[str] = None) -> str:
    """요약을 txt 파일로 저장 (백업용)"""
    if output_dir is None:
        output_dir = OUTPUT_DIR

    os.makedirs(output_dir, exist_ok=True)

    date_str = datetime.now().strftime('%Y%m%d')
    filename = f"{company}_{date_str}_{quarter}_컨콜요약.txt"
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(summary)

    return filepath


# =============================================================================
# 텔레그램 발송
# =============================================================================

def send_to_telegram(summary: str, company: str, quarter: str) -> bool:
    """요약을 텔레그램으로 발송"""
    if not BOT_TOKEN or not CHAT_ID:
        print("[경고] 텔레그램 설정이 없습니다. (BOT_TOKEN, CHAT_ID)")
        return False

    import requests

    # 메시지 구성 (텔레그램 메시지 길이 제한: 4096자)
    header = f"📊 *{company} {quarter} 컨콜 요약*\n\n"

    # 요약이 너무 길면 핵심만 추출
    if len(summary) > 3500:
        # Comment 섹션까지만 발송
        lines = summary.split('\n')
        short_summary = []
        for line in lines:
            short_summary.append(line)
            if line.startswith('* Comment') or line.startswith('Q&A'):
                # Q&A 시작 전까지 + Comment 섹션 일부
                if line.startswith('Q&A'):
                    short_summary.append("(Q&A는 파일 참조)")
                    break
        summary = '\n'.join(short_summary)

    message = header + summary

    # 텔레그램 전송
    url = f"https://api.telegram.org/bot{BOT_TOKEN.strip()}/sendMessage"

    try:
        response = requests.post(url, data={
            'chat_id': CHAT_ID.strip(),
            'text': message[:4096],
            'parse_mode': 'Markdown'
        }, timeout=30)

        if response.status_code == 200:
            print("[텔레그램] 발송 완료")
            return True
        else:
            # Markdown 파싱 실패시 일반 텍스트로 재시도
            response = requests.post(url, data={
                'chat_id': CHAT_ID.strip(),
                'text': message[:4096]
            }, timeout=30)
            if response.status_code == 200:
                print("[텔레그램] 발송 완료 (일반 텍스트)")
                return True
            print(f"[텔레그램] 발송 실패: {response.text}")
            return False

    except Exception as e:
        print(f"[텔레그램] 발송 오류: {e}")
        return False


# =============================================================================
# 대화형 모드
# =============================================================================

def interactive_mode():
    """대화형 모드로 컨콜 원문 입력받기"""
    print("\n" + "=" * 60)
    print("  컨퍼런스콜 원문 입력 (Ctrl+Z + Enter로 종료)")
    print("=" * 60)
    print()

    lines = []
    try:
        while True:
            line = input()
            lines.append(line)
    except EOFError:
        pass

    return '\n'.join(lines)


# =============================================================================
# 메인 실행
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='실적발표 컨퍼런스콜 요약')
    parser.add_argument('--file', '-f', type=str, help='컨콜 원문 파일 경로 (.txt, .docx)')
    parser.add_argument('--interactive', '-i', action='store_true', help='대화형 모드 (직접 입력)')
    parser.add_argument('--telegram', '-t', action='store_true', help='텔레그램 발송')
    parser.add_argument('--model', '-m', type=str, default='gpt-4o', help='GPT 모델 (기본: gpt-4o)')

    args = parser.parse_args()

    print("=" * 60)
    print("  실적발표 컨퍼런스콜 요약기")
    print(f"  실행 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 60)

    # 원문 가져오기
    if args.file:
        print(f"\n[1/3] 파일 읽기: {args.file}")
        transcript = read_transcript_file(args.file)
        print(f"  -> {len(transcript):,}자 로드 완료")
    elif args.interactive:
        transcript = interactive_mode()
        if not transcript.strip():
            print("[오류] 입력된 내용이 없습니다.")
            return
    else:
        print("\n사용법:")
        print("  python 4_Earnings_Call_Summarizer.py --file=원문파일.txt")
        print("  python 4_Earnings_Call_Summarizer.py --file=원문파일.docx --telegram")
        print("  python 4_Earnings_Call_Summarizer.py --interactive")
        return

    # GPT 요약
    print("\n[2/3] GPT 요약 생성 중...")
    try:
        summary, company, quarter = summarize_with_gpt(transcript, model=args.model)
        print(f"  -> 회사: {company}, 분기: {quarter}")
        print(f"  -> 요약 생성 완료 ({len(summary):,}자)")
    except Exception as e:
        print(f"[오류] GPT 요약 실패: {e}")
        return

    # 저장
    print("\n[3/3] 파일 저장 중...")
    try:
        docx_path = save_to_docx(summary, company, quarter)
        txt_path = save_to_txt(summary, company, quarter)
    except Exception as e:
        print(f"[경고] docx 저장 실패: {e}")
        txt_path = save_to_txt(summary, company, quarter)
        print(f"[저장 완료] {txt_path} (txt)")

    # 텔레그램 발송
    if args.telegram:
        print("\n[추가] 텔레그램 발송...")
        send_to_telegram(summary, company, quarter)

    # 결과 출력
    print("\n" + "=" * 60)
    print("[요약 결과 미리보기]")
    print("=" * 60)
    # 처음 50줄만 출력
    preview_lines = summary.split('\n')[:50]
    for line in preview_lines:
        print(line)
    if len(summary.split('\n')) > 50:
        print("... (이하 생략)")

    print("\n" + "=" * 60)
    print("[완료] 모든 작업이 완료되었습니다.")
    print("=" * 60)


if __name__ == "__main__":
    main()
